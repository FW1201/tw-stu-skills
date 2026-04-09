#!/usr/bin/env python3
"""
tw_edu_doc_utils.py
臺灣教育 Skills 共用 Word 文件樣式工具庫
所有 skills 的 scripts/ 均可 import 本模組
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# ── 品牌色彩 ─────────────────────────────────────────
BLUE_DEEP   = RGBColor(0x1A, 0x52, 0x76)
BLUE_MID    = RGBColor(0x24, 0x71, 0xA3)
BLUE_LIGHT  = RGBColor(0xEB, 0xF5, 0xFB)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT  = RGBColor(0xF8, 0xF9, 0xFA)
DARK_TEXT   = RGBColor(0x1C, 0x2A, 0x35)
GREEN       = RGBColor(0x1E, 0x84, 0x49)
GOLD        = RGBColor(0xD4, 0xAC, 0x0D)
ORANGE      = RGBColor(0xCA, 0x6F, 0x1E)
RED_SOFT    = RGBColor(0xA9, 0x3A, 0x26)

def rgb_hex(c):
    return '{:02X}{:02X}{:02X}'.format(c[0], c[1], c[2])

# ── 基本元素 ──────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    tcPr.append(shd)

def set_cell_border(cell, color='2471A3', size='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_east_asia_font(run, font='標楷體'):
    rPr = run._r.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:eastAsia'), font)
    rF.set(qn('w:ascii'), 'Arial')
    rPr.insert(0, rF)

def cell_write(cell, text, bold=False, size=11,
               color=DARK_TEXT, center=False, font='標楷體'):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = font
    r.font.color.rgb = color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_east_asia_font(r, font)
    return p

def header_cell(cell, text, bg=BLUE_MID, size=11):
    set_cell_bg(cell, bg)
    set_cell_border(cell, color=rgb_hex(BLUE_DEEP), size='6')
    cell_write(cell, text, bold=True, size=size,
               color=WHITE, center=True)

def data_cell(cell, text, row_idx=0, center=False):
    bg = BLUE_LIGHT if row_idx % 2 == 0 else GRAY_LIGHT
    set_cell_bg(cell, bg)
    set_cell_border(cell, color=rgb_hex(BLUE_MID))
    cell_write(cell, text, center=center)

def section_heading(doc, text, level=1):
    """帶底線的章節標題"""
    p = doc.add_paragraph()
    p.clear()
    prefix = '▌' if level == 1 else '▸'
    r = p.add_run(f'{prefix} {text}')
    r.bold = True
    r.font.size = Pt(14 if level == 1 else 12)
    r.font.name = '標楷體'
    r.font.color.rgb = BLUE_DEEP
    set_east_asia_font(r)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:color'), rgb_hex(BLUE_MID))
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def cover_page(doc, title, subtitle, info_pairs, accent_color=BLUE_DEEP):
    """通用封面頁"""
    doc.add_paragraph()
    doc.add_paragraph()

    # 主標題
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(26)
    r.font.name = '標楷體'; r.font.color.rgb = accent_color
    set_east_asia_font(r)

    # 副標題
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(16)
        r2.font.name = '標楷體'; r2.font.color.rgb = BLUE_MID
        set_east_asia_font(r2)

    doc.add_paragraph()

    # 資訊表（2欄成對）
    if info_pairs:
        tbl = doc.add_table(rows=(len(info_pairs)+1)//2, cols=4)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'
        col_widths = [Cm(3), Cm(5.5), Cm(3), Cm(5.5)]
        for col_idx, w in enumerate(col_widths):
            tbl.columns[col_idx].width = w

        flat = list(info_pairs.items())
        for i in range(0, len(flat), 2):
            row = tbl.rows[i // 2]
            k1, v1 = flat[i]
            header_cell(row.cells[0], k1)
            data_cell(row.cells[1], v1, row_idx=i//2)
            if i + 1 < len(flat):
                k2, v2 = flat[i+1]
                header_cell(row.cells[2], k2)
                data_cell(row.cells[3], v2, row_idx=i//2)
            else:
                data_cell(row.cells[2], '', row_idx=i//2)
                data_cell(row.cells[3], '', row_idx=i//2)

    doc.add_paragraph()
    doc.add_page_break()

def new_doc_a4(default_font='標楷體'):
    """建立標準 A4 文件"""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = default_font
    style.font.size = Pt(11)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), default_font)
    except Exception:
        pass
    return doc

def add_header_footer(doc, header_text, show_page=True):
    """加入頁首與頁碼"""
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.text = header_text
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if hp.runs:
        hp.runs[0].font.size = Pt(9)
        hp.runs[0].font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    if show_page:
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        for tag, text in [('begin', None), (None, ' PAGE '), ('end', None)]:
            if tag:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), tag)
                run._r.append(fc)
            else:
                it = OxmlElement('w:instrText')
                it.text = text
                run._r.append(it)
